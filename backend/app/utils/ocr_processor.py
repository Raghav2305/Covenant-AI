"""
OCR and Document Processing Utilities
"""

import asyncio
import os
from typing import Optional
import structlog
from docx import Document
import PyPDF2
import pytesseract
from PIL import Image
import pdf2image

logger = structlog.get_logger()


class OCRProcessor:
    """Document text extraction service"""
    
    def __init__(self):
        # Configure Tesseract path if needed
        # pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        pass
    
    async def extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            # First try direct text extraction
            text = await self._extract_pdf_text(file_path)
            
            # If no text found, try OCR
            if not text or len(text.strip()) < 100:
                logger.info("PDF text extraction yielded minimal results, trying OCR", file_path=file_path)
                text = await self._extract_pdf_with_ocr(file_path)
            
            return text
            
        except Exception as e:
            logger.error("PDF text extraction failed", file_path=file_path, error=str(e))
            raise
    
    async def extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            text = "\n".join(text_parts)
            logger.info("DOCX text extraction completed", 
                       file_path=file_path, 
                       text_length=len(text))
            
            return text
            
        except Exception as e:
            logger.error("DOCX text extraction failed", file_path=file_path, error=str(e))
            raise
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text directly from PDF"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_parts = []
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(page_text.strip())
                
                text = "\n".join(text_parts)
                logger.info("PDF direct text extraction completed", 
                           file_path=file_path, 
                           pages=len(pdf_reader.pages),
                           text_length=len(text))
                
                return text
                
        except Exception as e:
            logger.error("PDF direct text extraction failed", file_path=file_path, error=str(e))
            return ""
    
    async def _extract_pdf_with_ocr(self, file_path: str) -> str:
        """Extract text from PDF using OCR"""
        try:
            # Convert PDF to images
            images = pdf2image.convert_from_path(file_path)
            text_parts = []
            
            for i, image in enumerate(images):
                try:
                    # Perform OCR on each page
                    page_text = pytesseract.image_to_string(image, lang='eng')
                    if page_text.strip():
                        text_parts.append(f"--- Page {i+1} ---\n{page_text.strip()}")
                except Exception as e:
                    logger.warning("OCR failed for page", page=i+1, error=str(e))
                    continue
            
            text = "\n\n".join(text_parts)
            logger.info("PDF OCR extraction completed", 
                       file_path=file_path, 
                       pages=len(images),
                       text_length=len(text))
            
            return text
            
        except Exception as e:
            logger.error("PDF OCR extraction failed", file_path=file_path, error=str(e))
            raise
    
    async def extract_from_image(self, file_path: str) -> str:
        """Extract text from image file using OCR"""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image, lang='eng')
            
            logger.info("Image OCR extraction completed", 
                       file_path=file_path, 
                       text_length=len(text))
            
            return text.strip()
            
        except Exception as e:
            logger.error("Image OCR extraction failed", file_path=file_path, error=str(e))
            raise
    
    def validate_file(self, file_path: str) -> bool:
        """Validate if file can be processed"""
        if not os.path.exists(file_path):
            return False
        
        # Check file size (max 50MB)
        file_size = os.path.getsize(file_path)
        max_size = 50 * 1024 * 1024  # 50MB
        if file_size > max_size:
            logger.warning("File too large", file_path=file_path, size=file_size)
            return False
        
        # Check file extension
        allowed_extensions = ['.pdf', '.docx', '.doc', '.txt', '.png', '.jpg', '.jpeg']
        file_ext = os.path.splitext(file_path)[1].lower()
        if file_ext not in allowed_extensions:
            logger.warning("Unsupported file type", file_path=file_path, extension=file_ext)
            return False
        
        return True
    
    async def get_document_metadata(self, file_path: str) -> dict:
        """Extract document metadata"""
        try:
            stat = os.stat(file_path)
            return {
                "file_name": os.path.basename(file_path),
                "file_size": stat.st_size,
                "file_extension": os.path.splitext(file_path)[1].lower(),
                "created_at": stat.st_ctime,
                "modified_at": stat.st_mtime
            }
        except Exception as e:
            logger.error("Failed to get document metadata", file_path=file_path, error=str(e))
            return {}
    
    async def preprocess_image(self, image_path: str) -> str:
        """Preprocess image for better OCR results"""
        try:
            image = Image.open(image_path)
            
            # Convert to grayscale
            if image.mode != 'L':
                image = image.convert('L')
            
            # Enhance contrast
            from PIL import ImageEnhance
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(2.0)
            
            # Save processed image
            processed_path = image_path.replace('.', '_processed.')
            image.save(processed_path)
            
            return processed_path
            
        except Exception as e:
            logger.error("Image preprocessing failed", image_path=image_path, error=str(e))
            return image_path
